"""
Export Service for YouTube Summaries
Handles PDF, DOCX, and Markdown export functionality
"""
import io
import markdown
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
import logging

logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()

    def _strip_markdown(self, text: str) -> str:
        """Strip markdown formatting from text"""
        if not text:
            return ""
        import re
        # Remove bold/italic ** or *
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Remove headers #
        text = re.sub(r'#{1,6}\s?', '', text)
        # Remove code blocks
        text = re.sub(r'`{3}.*?`{3}', '', text, flags=re.DOTALL)
        text = re.sub(r'`(.*?)`', r'\1', text)
        # Remove links [text](url) -> text
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        return text.strip()
        
    def export_to_markdown(self, session_data: Dict[str, Any]) -> str:
        """Export session to Markdown format"""
        try:
            markdown_content = f"""# YouTube Video Summary
---

## Video Information
- **Title:** {session_data.get('video_title', 'Unknown')}
- **URL:** {session_data.get('video_url', '')}
- **Duration:** {self._format_duration(session_data.get('video_duration', 0))}
- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

## Short Summary
{session_data.get('short_summary', 'No summary available')}

---

## Detailed Summary
{session_data.get('detailed_summary', 'No detailed summary available')}

---

## Transcript
<details>
<summary>Click to view transcript</summary>

{session_data.get('transcript', 'No transcript available')}
</details>

---

## Chat History
"""
            
            chat_history = session_data.get('chat_history', [])
            if chat_history:
                for i, message in enumerate(chat_history):
                    role = message.get('role', 'unknown').title()
                    content = message.get('content', '')
                    timestamp = message.get('timestamp', '')
                    
                    if role.lower() == 'user':
                        markdown_content += f"\n### 👤 You ({timestamp})\n{content}\n"
                    else:
                        markdown_content += f"\n### 🤖 Assistant ({timestamp})\n> {content.replace(chr(10), chr(10) + '> ')}\n"
            else:
                markdown_content += "\n*No chat history available.*\n"
            
            return markdown_content
            
        except Exception as e:
            logger.error(f"Error exporting to markdown: {e}")
            raise
    
    def export_to_docx(self, session_data: Dict[str, Any]) -> io.BytesIO:
        """Export session to DOCX format"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('YouTube Video Summary', 0)
            title.alignment = 1  # Center alignment
            
            # Video Information
            doc.add_heading('Video Information', level=1)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ['Title', session_data.get('video_title', 'Unknown')],
                ['URL', session_data.get('video_url', '')],
                ['Duration', self._format_duration(session_data.get('video_duration', 0))],
                ['Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = value
            
            doc.add_page_break()
            
            # Short Summary
            doc.add_heading('Short Summary', level=1)
            doc.add_paragraph(self._strip_markdown(session_data.get('short_summary', 'No summary available')))
            
            # Detailed Summary
            doc.add_heading('Detailed Summary', level=1)
            doc.add_paragraph(self._strip_markdown(session_data.get('detailed_summary', 'No detailed summary available')))
            
            doc.add_page_break()
            
            # Chat History
            doc.add_heading('Chat History', level=1)
            chat_history = session_data.get('chat_history', [])
            
            if chat_history:
                for message in chat_history:
                    role = message.get('role', 'unknown').title()
                    content = self._strip_markdown(message.get('content', ''))
                    timestamp = message.get('timestamp', '')
                    
                    # Create role paragraph with distinct styling
                    role_para = doc.add_paragraph()
                    role_run = role_para.add_run(f"{role} ({timestamp})")
                    role_run.bold = True
                    
                    if role.lower() == 'assistant':
                        role_run.font.color.rgb = docx.shared.RGBColor(0x25, 0x63, 0xEB)  # Blue for AI
                    else:
                        role_run.font.color.rgb = docx.shared.RGBColor(0x05, 0x96, 0x69)  # Green for User
                    
                    # Add content
                    content_para = doc.add_paragraph(content)
                    
                    # Indent AI responses slightly to create a threaded look
                    if role.lower() == 'assistant':
                        content_para.paragraph_format.left_indent = Inches(0.25)
                        
                    doc.add_paragraph()  # Add space
            else:
                doc.add_paragraph('No chat history available.')
            
            doc.add_page_break()
            
            # Transcript
            doc.add_heading('Full Transcript', level=1)
            transcript_para = doc.add_paragraph(session_data.get('transcript', 'No transcript available'))
            transcript_para.style = 'Normal'
            
            # Save to BytesIO
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise
    
    def export_to_pdf(self, session_data: Dict[str, Any]) -> io.BytesIO:
        """Export session to PDF format"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
            
            # Define styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=self.styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1,  # Center
                textColor=HexColor('#2563eb')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=self.styles['Heading2'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=12,
                textColor=HexColor('#1f2937')
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leading=14
            )
            
            chat_role_style = ParagraphStyle(
                'ChatRole',
                parent=self.styles['Normal'],
                fontSize=12,
                spaceBefore=10,
                spaceAfter=5,
                textColor=HexColor('#059669')
            )
            
            # Build content
            content = []
            
            # Title
            content.append(Paragraph("YouTube Video Summary", title_style))
            content.append(Spacer(1, 20))
            
            # Video Information
            content.append(Paragraph("Video Information", heading_style))
            
            video_info = f"""
            <b>Title:</b> {session_data.get('video_title', 'Unknown')}<br/>
            <b>URL:</b> {session_data.get('video_url', '')}<br/>
            <b>Duration:</b> {self._format_duration(session_data.get('video_duration', 0))}<br/>
            <b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            """
            content.append(Paragraph(video_info, body_style))
            content.append(Spacer(1, 20))
            
            # Short Summary
            content.append(Paragraph("Short Summary", heading_style))
            content.append(Paragraph(self._strip_markdown(session_data.get('short_summary', 'No summary available')), body_style))
            content.append(Spacer(1, 20))
            
            # Detailed Summary
            content.append(Paragraph("Detailed Summary", heading_style))
            content.append(Paragraph(self._strip_markdown(session_data.get('detailed_summary', 'No detailed summary available')), body_style))
            content.append(PageBreak())
            
            # Chat History
            content.append(Paragraph("Chat History", heading_style))
            chat_history = session_data.get('chat_history', [])
            
            user_style = ParagraphStyle(
                'UserRole',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceBefore=10,
                spaceAfter=2,
                textColor=HexColor('#059669'), # Green
                fontName='Helvetica-Bold'
            )
            
            assistant_style = ParagraphStyle(
                'AssistantRole',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceBefore=10,
                spaceAfter=2,
                textColor=HexColor('#2563eb'), # Blue
                fontName='Helvetica-Bold'
            )
            
            message_style = ParagraphStyle(
                'MessageBody',
                parent=self.styles['Normal'],
                fontSize=10,
                spaceAfter=8,
                leading=14,
                leftIndent=10
            )
            
            if chat_history:
                for message in chat_history:
                    role = message.get('role', 'unknown').title()
                    content_text = self._strip_markdown(message.get('content', ''))
                    timestamp = message.get('timestamp', '')
                    
                    if role.lower() == 'user':
                        content.append(Paragraph(f"👤 {role} ({timestamp})", user_style))
                    else:
                        content.append(Paragraph(f"🤖 {role} ({timestamp})", assistant_style))
                        
                    content.append(Paragraph(content_text, message_style))
            else:
                content.append(Paragraph('No chat history available.', body_style))
            
            content.append(PageBreak())
            
            # Transcript
            content.append(Paragraph("Full Transcript", heading_style))
            transcript_text = session_data.get('transcript', 'No transcript available')
            # Split long transcript into paragraphs for better PDF formatting
            transcript_paragraphs = transcript_text.split('\n\n')
            for para in transcript_paragraphs:
                if para.strip():
                    content.append(Paragraph(para, body_style))
            
            # Build PDF
            doc.build(content)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    def export_document_to_pdf(self, document: dict) -> bytes:
        """
        Export a notes Document (with HTML content) to PDF bytes.
        Parses the stored HTML and maps tags to ReportLab styles.
        """
        try:
            from bs4 import BeautifulSoup
            buffer = io.BytesIO()
            doc_pdf = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                topMargin=1 * inch,
                leftMargin=1 * inch,
                rightMargin=1 * inch,
                bottomMargin=1 * inch,
            )

            title_style = ParagraphStyle(
                'DocTitle',
                parent=self.styles['Heading1'],
                fontSize=22,
                spaceAfter=20,
                textColor=HexColor('#2563eb'),
                alignment=0,
            )
            h1_style = ParagraphStyle(
                'DocH1',
                parent=self.styles['Heading1'],
                fontSize=18,
                spaceBefore=16,
                spaceAfter=10,
                textColor=HexColor('#1f2937'),
            )
            h2_style = ParagraphStyle(
                'DocH2',
                parent=self.styles['Heading2'],
                fontSize=15,
                spaceBefore=14,
                spaceAfter=8,
                textColor=HexColor('#374151'),
            )
            h3_style = ParagraphStyle(
                'DocH3',
                parent=self.styles['Heading3'],
                fontSize=13,
                spaceBefore=12,
                spaceAfter=6,
                textColor=HexColor('#4b5563'),
            )
            body_style = ParagraphStyle(
                'DocBody',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceAfter=8,
                leading=16,
                textColor=HexColor('#111827'),
            )
            meta_style = ParagraphStyle(
                'DocMeta',
                parent=self.styles['Normal'],
                fontSize=9,
                spaceAfter=20,
                textColor=HexColor('#6b7280'),
            )

            content_parts = []

            # Document title
            doc_title = document.get('title', 'Untitled Document')
            content_parts.append(Paragraph(doc_title, title_style))

            # Meta line
            updated = document.get('updated_at', '')
            if updated:
                try:
                    if hasattr(updated, 'strftime'):
                        updated_str = updated.strftime('%Y-%m-%d %H:%M')
                    else:
                        updated_str = str(updated)[:16]
                except Exception:
                    updated_str = str(updated)
                content_parts.append(Paragraph(f"Last updated: {updated_str}", meta_style))

            content_parts.append(Spacer(1, 12))

            # Parse HTML content
            raw_html = document.get('content', '') or ''
            soup = BeautifulSoup(raw_html, 'html.parser')

            def escape_xml(text):
                """Escape special chars for ReportLab XML."""
                return (text
                        .replace('&', '&amp;')
                        .replace('<', '&lt;')
                        .replace('>', '&gt;')
                        .replace('"', '&quot;'))

            def node_to_rl_text(node):
                """Recursively convert an HTML node to a ReportLab XML string."""
                if isinstance(node, str):
                    return escape_xml(node)
                tag = node.name.lower() if node.name else ''
                inner = ''.join(node_to_rl_text(c) for c in node.children)
                if tag in ('b', 'strong'):
                    return f'<b>{inner}</b>'
                elif tag in ('i', 'em'):
                    return f'<i>{inner}</i>'
                elif tag == 'u':
                    return f'<u>{inner}</u>'
                elif tag in ('mark',):
                    return inner  # ignore highlight colour in PDF
                elif tag == 'br':
                    return '<br/>'
                elif tag == 'span':
                    style_attr = node.get('style', '')
                    if 'font-size' in style_attr:
                        # best-effort inline font size
                        import re
                        m = re.search(r'font-size:\s*([\d.]+)px', style_attr)
                        if m:
                            size = max(8, int(float(m.group(1))))
                            return f'<font size="{size}">{inner}</font>'
                    return inner
                else:
                    return inner

            for element in soup.children:
                if isinstance(element, str):
                    text = element.strip()
                    if text:
                        content_parts.append(Paragraph(escape_xml(text), body_style))
                    continue

                tag = element.name.lower() if element.name else ''

                if tag == 'h1':
                    content_parts.append(Paragraph(element.get_text(), h1_style))
                elif tag == 'h2':
                    content_parts.append(Paragraph(element.get_text(), h2_style))
                elif tag in ('h3', 'h4', 'h5', 'h6'):
                    content_parts.append(Paragraph(element.get_text(), h3_style))
                elif tag in ('p', 'div'):
                    inner_xml = ''.join(node_to_rl_text(c) for c in element.children)
                    if inner_xml.strip():
                        try:
                            content_parts.append(Paragraph(inner_xml, body_style))
                        except Exception:
                            content_parts.append(Paragraph(escape_xml(element.get_text()), body_style))
                elif tag in ('ul', 'ol'):
                    for li in element.find_all('li', recursive=False):
                        bullet_text = f'• {escape_xml(li.get_text())}'
                        content_parts.append(Paragraph(bullet_text, body_style))
                elif tag == 'br':
                    content_parts.append(Spacer(1, 6))
                else:
                    # Generic fallback
                    text = element.get_text().strip()
                    if text:
                        content_parts.append(Paragraph(escape_xml(text), body_style))

            doc_pdf.build(content_parts)
            buffer.seek(0)
            return buffer.read()

        except Exception as e:
            logger.error(f"Error exporting document to PDF: {e}")
            raise

    def _format_duration(self, seconds: int) -> str:
        """Format duration in seconds to readable format"""
        if seconds <= 0:
            return "Unknown"
        
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        seconds = seconds % 60
        
        if hours > 0:
            return f"{hours}h {minutes}m {seconds}s"
        elif minutes > 0:
            return f"{minutes}m {seconds}s"
        else:
            return f"{seconds}s"

# Global instance
export_service = ExportService()
