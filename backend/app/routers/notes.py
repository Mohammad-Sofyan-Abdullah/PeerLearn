"""
Notes API Routes for Document Management and AI-Assisted Note Taking
"""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Body, BackgroundTasks
from typing import List, Optional
from datetime import datetime
from app.models import (
    Document, DocumentCreate, DocumentUpdate, DocumentChatMessage,
    UserInDB, DocumentStatus, DocumentSession, DocumentSessionCreate,
    DocumentSessionChatMessage, Flashcard, Quiz, QuizQuestion
)
from app.auth import get_current_active_user
from app.database import get_database
from app.ai_service import ai_service
from app.ocr_service import ocr_service
from bson import ObjectId
import logging
import os
import shutil
import docx
import PyPDF2
from io import BytesIO

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/notes", tags=["notes"])

# Upload directory for documents
UPLOAD_DIR = "static/uploads/documents"
os.makedirs(UPLOAD_DIR, exist_ok=True)


@router.get("/documents", response_model=List[Document])
async def get_user_documents(
    search: Optional[str] = None,
    status: Optional[DocumentStatus] = None,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get all documents for the current user"""
    try:
        # Build query
        query = {"user_id": ObjectId(current_user.id)}
        
        if status:
            query["status"] = status
        
        # Get documents
        cursor = db.documents.find(query).sort("updated_at", -1)
        documents = await cursor.to_list(length=None)
        
        # Convert ObjectIds to strings and add id field
        for doc in documents:
            doc["id"] = str(doc["_id"])
        
        # Filter by search if provided
        if search:
            search_lower = search.lower()
            documents = [
                doc for doc in documents 
                if search_lower in doc.get("title", "").lower() or 
                   search_lower in doc.get("content", "").lower()
            ]
        
        return [Document(**doc) for doc in documents]
        
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch documents"
        )

@router.post("/documents", response_model=Document)
async def create_document(
    document: DocumentCreate,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Create a new document"""
    try:
        document_dict = {
            "user_id": ObjectId(current_user.id),
            "title": document.title,
            "content": document.content,
            "status": DocumentStatus.DRAFT,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.documents.insert_one(document_dict)
        document_dict["_id"] = result.inserted_id
        
        # Convert ObjectId to string for response
        document_dict["id"] = str(result.inserted_id)
        
        logger.info(f"Created document {result.inserted_id} for user {current_user.id}")
        return Document(**document_dict)
        
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create document"
        )

@router.get("/debug/test")
async def debug_test(
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Debug endpoint to test document creation and retrieval"""
    try:
        # Create a test document
        test_doc = {
            "user_id": ObjectId(current_user.id),
            "title": "Debug Test Document",
            "content": "This is a test document for debugging",
            "status": DocumentStatus.DRAFT,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.documents.insert_one(test_doc)
        test_doc["_id"] = result.inserted_id
        test_doc["id"] = str(result.inserted_id)
        
        logger.info(f"Created test document: {test_doc}")
        
        # Try to retrieve it
        retrieved = await db.documents.find_one({"_id": result.inserted_id})
        retrieved["id"] = str(retrieved["_id"])
        
        logger.info(f"Retrieved test document: {retrieved}")
        
        # Clean up
        await db.documents.delete_one({"_id": result.inserted_id})
        
        return {
            "created": test_doc,
            "retrieved": retrieved,
            "id_string": str(result.inserted_id),
            "id_valid": ObjectId.is_valid(str(result.inserted_id))
        }
        
    except Exception as e:
        logger.error(f"Debug test failed: {e}")
        return {"error": str(e)}

@router.get("/documents/{document_id}", response_model=Document)
async def get_document(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get a specific document"""
    logger.info(f"Getting document with ID: {document_id}")
    
    try:
        document_object_id = ObjectId(document_id)
        logger.info(f"Converted to ObjectId: {document_object_id}")
    except Exception as e:
        logger.error(f"Invalid document ID format: {document_id}, error: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format"
        )
    
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        logger.warning(f"Document not found: {document_id} for user {current_user.id}")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Convert ObjectId to string for response
    document["id"] = str(document["_id"])
    logger.info(f"Returning document: {document['title']}")
    
    return Document(**document)

@router.put("/documents/{document_id}", response_model=Document)
async def update_document(
    document_id: str,
    document_update: DocumentUpdate,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Update a document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Check if document exists and belongs to user
    existing_doc = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not existing_doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Build update data
    update_data = {"updated_at": datetime.utcnow()}
    if document_update.title is not None:
        update_data["title"] = document_update.title
    if document_update.content is not None:
        update_data["content"] = document_update.content
    if document_update.status is not None:
        update_data["status"] = document_update.status
    
    # Update document
    await db.documents.update_one(
        {"_id": document_object_id},
        {"$set": update_data}
    )
    
    # Return updated document
    updated_doc = await db.documents.find_one({"_id": document_object_id})
    updated_doc["id"] = str(updated_doc["_id"])
    return Document(**updated_doc)

@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Delete a document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Check if document exists and belongs to user
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Delete associated file if exists
    if document.get("file_url"):
        try:
            if os.path.exists(document["file_url"]):
                os.remove(document["file_url"])
        except Exception as e:
            logger.warning(f"Failed to delete file {document['file_url']}: {e}")
    
    # Delete document and chat history
    await db.documents.delete_one({"_id": document_object_id})
    await db.document_chat_messages.delete_many({"document_id": document_object_id})
    
    return {"message": "Document deleted successfully"}

@router.post("/documents/upload", response_model=Document)
async def upload_document(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Upload a document file and extract content (supports images, videos, PDFs, PowerPoints, and more)"""
    try:
        # Expanded list of allowed file types
        allowed_extensions = {
            # Text documents
            '.txt', '.docx', '.doc', '.pdf', '.rtf',
            # Presentations
            '.ppt', '.pptx', '.odp',
            # Images (for OCR)
            '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp', '.gif',
            # Videos (for OCR from frames)
            '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v',
            # Spreadsheets
            '.xls', '.xlsx', '.ods', '.csv'
        }
        
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type {file_extension} not supported. Supported types: {', '.join(sorted(allowed_extensions))}"
            )
        
        # Save file
        file_path = os.path.join(UPLOAD_DIR, f"{ObjectId()}_{file.filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract content based on file type
        content = ""
        extraction_method = "unknown"
        
        try:
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                extraction_method = "text_file"
                
            elif file_extension in ['.docx', '.doc']:
                doc = docx.Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                extraction_method = "docx_parser"
                
            elif file_extension == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = '\n'.join([page.extract_text() for page in pdf_reader.pages])
                extraction_method = "pdf_parser"
                
            elif file_extension in ['.ppt', '.pptx']:
                content = await ocr_service.extract_text_from_presentation(file_path)
                extraction_method = "presentation_parser"
                
            elif ocr_service.is_image_file(file.filename):
                logger.info(f"Processing image file with OCR: {file.filename}")
                content = await ocr_service.extract_text_from_image(file_path)
                extraction_method = "image_ocr"
                
            elif ocr_service.is_video_file(file.filename):
                logger.info(f"Processing video file with OCR: {file.filename}")
                content = await ocr_service.extract_text_from_video(file_path)
                extraction_method = "video_ocr"
                
            elif file_extension == '.csv':
                import pandas as pd
                df = pd.read_csv(file_path)
                content = df.to_string(index=False)
                extraction_method = "csv_parser"
                
            elif file_extension in ['.xls', '.xlsx']:
                import pandas as pd
                df = pd.read_excel(file_path)
                content = df.to_string(index=False)
                extraction_method = "excel_parser"
                
            else:
                # Try to read as text file as fallback
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    extraction_method = "text_fallback"
                except:
                    content = f"Content extraction not supported for {file_extension} files. You can edit this document manually."
                    extraction_method = "manual_edit_required"
                    
        except Exception as e:
            logger.warning(f"Failed to extract content from {file.filename}: {e}")
            if ocr_service.is_image_file(file.filename) or ocr_service.is_video_file(file.filename):
                content = f"OCR extraction failed for {file.filename}. Error: {str(e)}"
                extraction_method = "ocr_failed"
            else:
                content = f"Content extraction failed for {file.filename}. You can edit this document manually."
                extraction_method = "extraction_failed"
        
        # Use provided title or filename
        document_title = title or os.path.splitext(file.filename)[0]
        
        # Add extraction info to content if OCR was used
        if extraction_method in ["image_ocr", "video_ocr"]:
            if content and "extraction failed" not in content.lower():
                content = f"[Extracted using OCR from {extraction_method.replace('_', ' ')}]\n\n{content}"
            
        # Create document
        document_dict = {
            "user_id": ObjectId(current_user.id),
            "title": document_title,
            "content": content,
            "file_url": file_path,
            "file_name": file.filename,
            "file_size": os.path.getsize(file_path),
            "file_type": file_extension,
            "extraction_method": extraction_method,
            "status": DocumentStatus.DRAFT,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.documents.insert_one(document_dict)
        document_dict["_id"] = result.inserted_id
        
        # Convert ObjectId to string for response
        document_dict["id"] = str(result.inserted_id)
        
        logger.info(f"Uploaded document {result.inserted_id} for user {current_user.id} using {extraction_method}")
        return Document(**document_dict)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to upload document"
        )

@router.post("/documents/{document_id}/chat")
async def chat_with_document(
    document_id: str,
    message: str = Body(..., embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Chat with AI about the document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Get document
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        # Get recent chat history
        chat_history = await db.document_chat_messages.find({
            "document_id": document_object_id
        }).sort("timestamp", -1).limit(10).to_list(length=None)
        
        # Reverse to get chronological order
        chat_history.reverse()
        
        # Generate AI response
        ai_response = await ai_service.chat_with_document(
            content=document.get("content", ""),
            document_title=document.get("title", ""),
            user_message=message,
            chat_history=chat_history
        )
        
        # Save chat message
        chat_message = {
            "document_id": document_object_id,
            "user_id": ObjectId(current_user.id),
            "message": message,
            "response": ai_response,
            "timestamp": datetime.utcnow()
        }
        
        result = await db.document_chat_messages.insert_one(chat_message)
        chat_message["_id"] = result.inserted_id
        chat_message["id"] = str(result.inserted_id)
        
        return {
            "message": message,
            "response": ai_response,
            "timestamp": chat_message["timestamp"]
        }
        
    except Exception as e:
        logger.error(f"Error in document chat: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process chat message"
        )

@router.post("/documents/{document_id}/reprocess-ocr")
async def reprocess_document_with_ocr(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Reprocess a document using OCR if initial extraction failed"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Get document
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    file_path = document.get("file_url")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Original file not found"
        )
    
    try:
        filename = document.get("file_name", "")
        content = ""
        extraction_method = "reprocessed_unknown"
        
        if ocr_service.is_image_file(filename):
            logger.info(f"Reprocessing image file with OCR: {filename}")
            content = await ocr_service.extract_text_from_image(file_path)
            extraction_method = "reprocessed_image_ocr"
            
        elif ocr_service.is_video_file(filename):
            logger.info(f"Reprocessing video file with OCR: {filename}")
            content = await ocr_service.extract_text_from_video(file_path)
            extraction_method = "reprocessed_video_ocr"
            
        elif ocr_service.is_presentation_file(filename):
            logger.info(f"Reprocessing presentation file: {filename}")
            content = await ocr_service.extract_text_from_presentation(file_path)
            extraction_method = "reprocessed_presentation"
            
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File type not supported for OCR reprocessing"
            )
        
        # Add reprocessing info to content
        if content and "extraction failed" not in content.lower():
            content = f"[Reprocessed using OCR - {extraction_method.replace('_', ' ')}]\n\n{content}"
        
        # Update document
        await db.documents.update_one(
            {"_id": document_object_id},
            {
                "$set": {
                    "content": content,
                    "extraction_method": extraction_method,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {
            "message": "Document reprocessed successfully",
            "extraction_method": extraction_method,
            "content_length": len(content),
            "content_preview": content[:200] + "..." if len(content) > 200 else content
        }
        
    except Exception as e:
        logger.error(f"Error reprocessing document with OCR: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to reprocess document: {str(e)}"
        )
async def generate_notes(
    document_id: str,
    prompt: str = Body(..., embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Generate structured notes from document content"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Get document
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        # Generate notes using AI
        notes = await ai_service.generate_notes_from_document(
            content=document.get("content", ""),
            document_title=document.get("title", ""),
            user_prompt=prompt
        )
        
        return {
            "notes": notes,
            "prompt": prompt
        }
        
    except Exception as e:
        logger.error(f"Error generating notes: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate notes"
        )

@router.get("/documents/{document_id}/chat-history")
async def get_chat_history(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get chat history for a document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Verify document ownership
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Get chat history
    chat_history = await db.document_chat_messages.find({
        "document_id": document_object_id
    }).sort("timestamp", 1).to_list(length=None)
    
    # Convert ObjectIds to strings
    for chat in chat_history:
        chat["id"] = str(chat["_id"])
        chat["_id"] = str(chat["_id"])
        if "document_id" in chat and isinstance(chat["document_id"], ObjectId):
            chat["document_id"] = str(chat["document_id"])
        if "user_id" in chat and isinstance(chat["user_id"], ObjectId):
            chat["user_id"] = str(chat["user_id"])
    
    return {"chat_history": chat_history}


# ============================================
# DOCUMENT SESSION ENDPOINTS
# ============================================

@router.post("/sessions", response_model=DocumentSession)
async def create_document_session(
    session_data: DocumentSessionCreate,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Create a new document session from an uploaded document"""
    try:
        document_object_id = ObjectId(session_data.document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Get the document
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Create session document
    session_dict = {
        "user_id": ObjectId(current_user.id),
        "document_id": document_object_id,
        "document_title": document.get("title", "Untitled Document"),
        "document_content": document.get("content", ""),
        "short_summary": None,
        "detailed_summary": None,
        "chat_history": [],
        "flashcards": [],
        "quiz": None,
        "slides_pdf_url": None,
        "slides_status": "pending",
        "generated_slide_images": [],
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    result = await db.document_sessions.insert_one(session_dict)
    session_dict["_id"] = result.inserted_id
    
    logger.info(f"Created document session {result.inserted_id} for document {session_data.document_id}")
    return DocumentSession(**session_dict)


@router.get("/sessions", response_model=List[DocumentSession])
async def get_user_document_sessions(
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get all document sessions for the current user"""
    try:
        sessions = await db.document_sessions.find(
            {"user_id": ObjectId(current_user.id)}
        ).sort("created_at", -1).to_list(length=100)
        
        # Convert ObjectIds to strings
        for session in sessions:
            session["_id"] = str(session["_id"])
            session["user_id"] = str(session["user_id"])
            session["document_id"] = str(session["document_id"])
        
        return [DocumentSession(**session) for session in sessions]
        
    except Exception as e:
        logger.error(f"Error fetching document sessions: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch sessions"
        )


@router.get("/sessions/{session_id}", response_model=DocumentSession)
async def get_document_session(
    session_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get a specific document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    # Convert ObjectIds to strings
    session["_id"] = str(session["_id"])
    session["user_id"] = str(session["user_id"])
    session["document_id"] = str(session["document_id"])
    
    return DocumentSession(**session)


@router.delete("/sessions/{session_id}")
async def delete_document_session(
    session_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Delete a document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    result = await db.document_sessions.delete_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if result.deleted_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    return {"message": "Session deleted successfully"}


@router.post("/sessions/{session_id}/summarize")
async def summarize_document_session(
    session_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Generate summaries for a document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    if not session.get("document_content"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No document content available"
        )
    
    try:
        short_summary, detailed_summary = await ai_service.generate_document_summaries(
            content=session["document_content"],
            document_title=session["document_title"]
        )
        
        # Update session
        await db.document_sessions.update_one(
            {"_id": session_object_id},
            {
                "$set": {
                    "short_summary": short_summary,
                    "detailed_summary": detailed_summary,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {
            "short_summary": short_summary,
            "detailed_summary": detailed_summary
        }
        
    except Exception as e:
        logger.error(f"Error summarizing document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to summarize document: {str(e)}"
        )


@router.post("/sessions/{session_id}/chat")
async def chat_with_document_session(
    session_id: str,
    message: str = Body(..., embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Chat with AI about the document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    try:
        # Get recent chat history
        chat_history = session.get("chat_history", [])[-10:]
        
        # Generate AI response
        ai_response = await ai_service.chat_with_document(
            content=session.get("document_content", ""),
            document_title=session.get("document_title", ""),
            user_message=message,
            chat_history=[{"message": msg["content"], "response": ""} for msg in chat_history if msg["role"] == "user"]
        )
        
        # Create chat messages
        user_message = {"role": "user", "content": message, "timestamp": datetime.utcnow()}
        assistant_message = {"role": "assistant", "content": ai_response, "timestamp": datetime.utcnow()}
        
        # Update session with new chat messages
        await db.document_sessions.update_one(
            {"_id": session_object_id},
            {
                "$push": {
                    "chat_history": {
                        "$each": [user_message, assistant_message]
                    }
                },
                "$set": {"updated_at": datetime.utcnow()}
            }
        )
        
        return {
            "message": message,
            "response": ai_response,
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error in document session chat: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process chat message: {str(e)}"
        )


@router.post("/sessions/{session_id}/chat/clear")
async def clear_session_chat_history(
    session_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Clear chat history for a document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
        
    await db.document_sessions.update_one(
        {"_id": session_object_id},
        {"$set": {"chat_history": [], "updated_at": datetime.utcnow()}}
    )
    
    return {"message": "Chat history cleared"}


@router.post("/sessions/{session_id}/flashcards")
async def generate_document_flashcards(
    session_id: str,
    count: int = Body(15, embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Generate flashcards for a document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    # Use summary if available, otherwise use content
    content = session.get("detailed_summary") or session.get("document_content", "")
    
    if not content or len(content.strip()) < 50:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Not enough content to generate flashcards. Please summarize the document first."
        )
    
    try:
        flashcards_data = await ai_service.generate_document_flashcards(
            content=content,
            document_title=session["document_title"],
            count=count
        )
        
        # Convert to Flashcard models
        flashcards = [Flashcard(**card) for card in flashcards_data]
        
        # Update session
        await db.document_sessions.update_one(
            {"_id": session_object_id},
            {
                "$set": {
                    "flashcards": [card.dict() for card in flashcards],
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {
            "flashcards": flashcards,
            "count": len(flashcards)
        }
        
    except ValueError as e:
        return {
            "flashcards": [],
            "count": 0,
            "message": "Flashcards not available right now. Please try again later.",
            "error": str(e)
        }
    except Exception as e:
        logger.error(f"Error generating document flashcards: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred while generating flashcards"
        )


@router.post("/sessions/{session_id}/flashcards/explain")
async def explain_document_flashcard(
    session_id: str,
    question: str = Body(..., embed=True),
    answer: str = Body(..., embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get AI explanation for a specific flashcard"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    try:
        context = session.get("detailed_summary") or session.get("document_content", "")[:15000]
        
        explanation = await ai_service.explain_flashcard_answer(
            question=question,
            answer=answer,
            context=context,
            video_title=session["document_title"]  # Reusing the video method
        )
        
        return {
            "question": question,
            "answer": answer,
            "explanation": explanation
        }
        
    except Exception as e:
        logger.error(f"Error explaining flashcard: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate explanation: {str(e)}"
        )


@router.post("/sessions/{session_id}/quiz")
async def generate_document_quiz(
    session_id: str,
    count: int = Body(10, embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Generate a multiple-choice quiz for a document session"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    # Use summary if available, otherwise use content
    content = session.get("detailed_summary") or session.get("document_content", "")
    
    if not content or len(content.strip()) < 50:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Not enough content to generate quiz. Please summarize the document first."
        )
    
    try:
        quiz_data = await ai_service.generate_document_quiz(
            content=content,
            document_title=session["document_title"],
            count=count
        )
        
        # Set the generated_at timestamp
        quiz_data["generated_at"] = datetime.utcnow()
        
        # Update session
        await db.document_sessions.update_one(
            {"_id": session_object_id},
            {
                "$set": {
                    "quiz": quiz_data,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {
            "quiz": quiz_data,
            "count": len(quiz_data["questions"])
        }
        
    except ValueError as e:
        return {
            "quiz": None,
            "count": 0,
            "message": "Quiz not available right now. Please try again later.",
            "error": str(e)
        }
    except Exception as e:
        logger.error(f"Error generating document quiz: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred while generating quiz"
        )


async def process_document_slide_generation(session_id: str, summary: str, count: int, db):
    """Background task to generate slides for document session"""
    try:
        session_object_id = ObjectId(session_id)
        
        # Generate Content Plan
        slides_content = await ai_service.generate_slides_content(summary, count)
        
        image_urls = []
        output_dir = f"static/slides/document_{session_id}"
        os.makedirs(output_dir, exist_ok=True)

        # Generate Images
        for i, slide in enumerate(slides_content):
            if "image_prompt" in slide:
                try:
                    img_bytes = await ai_service.generate_slide_image(slide["image_prompt"])
                    
                    if img_bytes:
                        slide["image_bytes"] = img_bytes
                        img_filename = f"slide_{i+1}.png"
                        img_path = os.path.join(output_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(img_bytes)
                        
                        image_url = f"/static/slides/document_{session_id}/{img_filename}"
                        image_urls.append(image_url)
                        
                except Exception as e:
                    logger.error(f"Image generation failed for slide {slide.get('title')}: {e}")
                    slide["image_bytes"] = None
        
        # Create PDF
        pdf_filename = f"slides_{session_id}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        await ai_service.create_slides_pdf(slides_content, pdf_path)
        
        pdf_url = f"/static/slides/document_{session_id}/{pdf_filename}"
        
        # Update session
        await db.document_sessions.update_one(
            {"_id": session_object_id},
            {
                "$set": {
                    "slides_pdf_url": pdf_url,
                    "generated_slide_images": image_urls,
                    "slides_status": "completed"
                }
            }
        )
        logger.info(f"Slide generation completed for document session {session_id}")
        
    except Exception as e:
        logger.error(f"Error in background slide generation for document: {e}")
        try:
            await db.document_sessions.update_one(
                {"_id": ObjectId(session_id)},
                {"$set": {"slides_status": "failed"}}
            )
        except:
            pass


@router.post("/sessions/{session_id}/slides")
async def generate_document_slides(
    session_id: str,
    background_tasks: BackgroundTasks,
    count: int = Body(5, embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Generate visual slides for the document summary (Background Task)"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    session = await db.document_sessions.find_one({
        "_id": session_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
        
    summary = session.get("detailed_summary") or session.get("short_summary")
    if not summary:
        raise HTTPException(status_code=400, detail="No summary available. Please summarize the document first.")
        
    # Update status to processing
    await db.document_sessions.update_one(
        {"_id": session_object_id},
        {"$set": {"slides_status": "processing"}}
    )
    
    # Add to background tasks
    background_tasks.add_task(process_document_slide_generation, session_id, summary, count, db)
    
    return {"status": "processing", "message": "Slide generation started in background"}


@router.post("/sessions/{session_id}/regenerate-summaries")
async def regenerate_document_summaries(
    session_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Regenerate summaries for a document session"""
    # Just call the summarize endpoint
    return await summarize_document_session(session_id, current_user, db)


@router.post("/sessions/{session_id}/import")
async def import_shared_document_session(
    session_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Import a shared document session to the current user's account"""
    try:
        session_object_id = ObjectId(session_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid session ID"
        )
    
    # Find the original session (any user's session)
    original_session = await db.document_sessions.find_one({
        "_id": session_object_id
    })
    
    if not original_session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Session not found"
        )
    
    # Check if user already owns this session
    if str(original_session.get("user_id")) == str(current_user.id):
        return {
            "message": "You already own this session",
            "session_id": session_id,
            "already_owned": True
        }
    
    # Check if user already imported this session
    existing_import = await db.document_sessions.find_one({
        "user_id": ObjectId(current_user.id),
        "imported_from": session_id
    })
    
    if existing_import:
        return {
            "message": "Session already imported",
            "session_id": str(existing_import["_id"]),
            "already_imported": True
        }
    
    # Clone the session for the current user
    new_session = {
        "user_id": ObjectId(current_user.id),
        "document_id": original_session.get("document_id"),
        "document_title": original_session.get("document_title"),
        "document_content": original_session.get("document_content"),
        "short_summary": original_session.get("short_summary"),
        "detailed_summary": original_session.get("detailed_summary"),
        "chat_history": [],  # Start with fresh chat history
        "flashcards": original_session.get("flashcards", []),
        "quiz": original_session.get("quiz"),
        "slides_pdf_url": original_session.get("slides_pdf_url"),
        "slides_status": original_session.get("slides_status"),
        "generated_slide_images": original_session.get("generated_slide_images", []),
        "imported_from": session_id,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    result = await db.document_sessions.insert_one(new_session)
    
    return {
        "message": "Session imported successfully",
        "session_id": str(result.inserted_id),
        "imported": True
    }
