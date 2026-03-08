"""
DataMinors_Edu Project Documentation Generator

This script generates comprehensive Word documentation for the DataMinors_Edu project.
It analyzes the project structure and creates detailed documentation covering all aspects.

Usage:
    python generate_documentation.py

Output:
    DataMinors_Edu_Documentation.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a formatted heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph to the document"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    return p

def add_code_block(doc, code, language=""):
    """Add a code block with monospace font"""
    p = doc.add_paragraph(code)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)
    return p

def create_title_page(doc):
    """Create the title page"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DataMinors_Edu\n")
    run.font.size = Pt(36)
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Project Documentation")
    run.font.size = Pt(24)
    
    doc.add_paragraph()  # spacing
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents placeholder"""
    add_heading(doc, "Table of Contents", level=1)
    add_paragraph(doc, "1. Project Overview")
    add_paragraph(doc, "2. System Architecture")
    add_paragraph(doc, "3. Technology Stack")
    add_paragraph(doc, "4. Features")
    add_paragraph(doc, "5. Backend Documentation")
    add_paragraph(doc, "6. Frontend Documentation")
    add_paragraph(doc, "7. Database Schema")
    add_paragraph(doc, "8. API Documentation")
    add_paragraph(doc, "9. Setup & Installation")
    add_paragraph(doc, "10. Deployment Guide")
    doc.add_page_break()

def add_project_overview(doc):
    """Add project overview section"""
    add_heading(doc, "1. Project Overview", level=1)
    
    add_heading(doc, "1.1 Introduction", level=2)
    add_paragraph(doc, 
        "DataMinors_Edu is a comprehensive educational platform designed to enhance student learning through AI-powered tools "
        "and collaborative features. The platform combines document management, video summarization, real-time collaboration, "
        "and intelligent study aids to create an all-in-one learning environment.")
    
    add_heading(doc, "1.2 Key Objectives", level=2)
    objectives = [
        "Provide AI-assisted document analysis and note-taking",
        "Enable collaborative learning through virtual classrooms",
        "Generate study materials from YouTube videos",
        "Facilitate peer-to-peer learning and discussion",
        "Offer intelligent OCR for scanned documents",
        "Create flashcards and quizzes automatically"
    ]
    for objective in objectives:
        add_paragraph(doc, f"• {objective}")
    
    add_heading(doc, "1.3 Target Users", level=2)
    add_paragraph(doc, "• Students seeking AI-powered study assistance")
    add_paragraph(doc, "• Educators managing classroom discussions")
    add_paragraph(doc, "• Study groups collaborating remotely")
    add_paragraph(doc, "• Learners using video content for education")
    
    doc.add_page_break()

def add_architecture(doc):
    """Add system architecture section"""
    add_heading(doc, "2. System Architecture", level=1)
    
    add_heading(doc, "2.1 High-Level Architecture", level=2)
    add_paragraph(doc, 
        "DataMinors_Edu follows a modern three-tier architecture:")
    
    add_paragraph(doc, "• Presentation Layer: React-based SPA")
    add_paragraph(doc, "• Application Layer: FastAPI backend")
    add_paragraph(doc, "• Data Layer: MongoDB database")
    
    add_heading(doc, "2.2 Component Diagram", level=2)
    components = {
        "Frontend": [
            "React 18 with React Router",
            "TailwindCSS for styling",
            "Socket.IO client for real-time features",
            "React Query for state management"
        ],
        "Backend": [
            "FastAPI framework",
            "MongoDB with Motor (async driver)",
            "Socket.IO server",
            "JWT authentication",
            "AI integrations (Groq, Google Gemini)"
        ],
        "External Services": [
            "Groq API for AI capabilities",
            "Google Gemini for image generation",
            "YouTube API for video processing"
        ]
    }
    
    for component, features in components.items():
        add_heading(doc, component, level=3)
        for feature in features:
            add_paragraph(doc, f"• {feature}")
    
    doc.add_page_break()

def add_technology_stack(doc):
    """Add technology stack section"""
    add_heading(doc, "3. Technology Stack", level=1)
    
    add_heading(doc, "3.1 Frontend Technologies", level=2)
    frontend_tech = [
        ("React", "18.x", "UI library"),
        ("React Router", "6.x", "Client-side routing"),
        ("TailwindCSS", "3.x", "Utility-first CSS framework"),
        ("Framer Motion", "Latest", "Animation library"),
        ("React Query", "3.x", "Server state management"),
        ("Socket.IO Client", "4.x", "Real-time communication"),
        ("Lucide React", "Latest", "Icon library"),
        ("React Hot Toast", "Latest", "Toast notifications")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'
    
    for pkg, ver, purpose in frontend_tech:
        row_cells = table.add_row().cells
        row_cells[0].text = pkg
        row_cells[1].text = ver
        row_cells[2].text = purpose
    
    doc.add_paragraph()  # spacing
    
    add_heading(doc, "3.2 Backend Technologies", level=2)
    backend_tech = [
        ("FastAPI", "Latest", "Web framework"),
        ("MongoDB", "Latest", "NoSQL database"),
        ("Motor", "Latest", "Async MongoDB driver"),
        ("Python-Jose", "Latest", "JWT tokens"),
        ("Passlib", "Latest", "Password hashing"),
        ("Groq", "Latest", "AI/LLM API"),
        ("Google GenAI", "Latest", "Image generation"),
        ("OpenCV", "Latest", "Image processing"),
        ("Scikit-Image", "Latest", "Advanced image processing"),
        ("YT-DLP", "Latest", "YouTube video downloader"),
        ("Python-DOCX", "Latest", "Word document generation"),
        ("ReportLab", "Latest", "PDF generation"),
        ("Python-PPTX", "Latest", "PowerPoint processing")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'
    
    for pkg, ver, purpose in backend_tech:
        row_cells = table.add_row().cells
        row_cells[0].text = pkg
        row_cells[1].text = ver
        row_cells[2].text = purpose
    
    doc.add_page_break()

def add_features(doc):
    """Add features section"""
    add_heading(doc, "4. Core Features", level=1)
    
    features = {
        "4.1 User Management": [
            "User registration and authentication",
            "JWT-based session management",
            "Password reset functionality",
            "Friend system with friend requests",
            "User profiles and settings"
        ],
        "4.2 Document Management": [
            "Upload documents (PDF, DOCX, TXT, images, videos)",
            "AI-powered OCR with document scanning",
            "Edge detection and perspective correction",
            "Shadow removal and contrast enhancement",
            "AI-formatted text output",
            "Document sessions for AI interaction",
            "Chat with AI about documents",
            "Generate notes from documents",
            "Real-time document editing"
        ],
        "4.3 YouTube Summarizer": [
            "Video URL input and processing",
            "Automatic transcript extraction",
            "AI-generated short summaries",
            "Detailed content summaries",
            "Flashcard generation from videos",
            "Quiz creation",
            "Related video suggestions",
            "Slide generation with AI images",
            "Export to PDF"
        ],
        "4.4 Virtual Classrooms": [
            "Create and manage classrooms",
            "Invite members via invite codes",
            "Create multiple channels per classroom",
            "Admin and member roles",
            "Real-time chat with Socket.IO",
            "Message editing and deletion",
            "Content moderation with AI",
            "Chat summarization",
            "Classroom settings management"
        ],
        "4.5 Collaborative Features": [
            "Real-time messaging",
            "Friend system",
            "Share content with friends",
            "Classroom collaboration",
            "Socket.IO for live updates"
        ]
    }
    
    for section, items in features.items():
        add_heading(doc, section, level=2)
        for item in items:
            add_paragraph(doc, f"• {item}")
    
    doc.add_page_break()

def add_backend_docs(doc):
    """Add backend documentation"""
    add_heading(doc, "5. Backend Documentation", level=1)
    
    add_heading(doc, "5.1 Project Structure", level=2)
    add_code_block(doc, """backend/
├── app/
│   ├── main.py              # FastAPI application entry point
│   ├── config.py            # Configuration and environment variables
│   ├── database.py          # MongoDB connection
│   ├── models.py            # Pydantic models
│   ├── auth.py              # Authentication utilities
│   ├── socket_manager.py    # Socket.IO manager
│   ├── ai_service.py        # AI integration service
│   ├── ocr_service.py       # OCR and image processing
│   ├── routers/
│   │   ├── auth.py          # Auth endpoints
│   │   ├── users.py         # User management
│   │   ├── friends.py       # Friend system
│   │   ├── notes.py         # Document management
│   │   ├── classrooms.py    # Classroom endpoints
│   │   ├── chat.py          # Chat endpoints
│   │   └── youtube.py       # YouTube processing
│   └── ...
└── requirements.txt         # Python dependencies""")
    
    add_heading(doc, "5.2 Key Services", level=2)
    
    add_heading(doc, "OCR Service (ocr_service.py)", level=3)
    add_paragraph(doc, "Advanced OCR with document scanning capabilities:")
    add_paragraph(doc, "• Document edge detection using Canny algorithm")
    add_paragraph(doc, "• Perspective transformation for skewed documents")
    add_paragraph(doc, "• Shadow removal with LAB color space")
    add_paragraph(doc, "• CLAHE contrast enhancement")
    add_paragraph(doc, "• Otsu's binarization for clean text")
    add_paragraph(doc, "• AI-powered text formatting and cleanup")
    
    add_heading(doc, "AI Service (ai_service.py)", level=3)
    add_paragraph(doc, "Comprehensive AI integration:")
    add_paragraph(doc, "• Content moderation")
    add_paragraph(doc, "• Chat summarization")
    add_paragraph(doc, "• Flashcard generation")
    add_paragraph(doc, "• Quiz creation")
    add_paragraph(doc, "• Related video suggestions")
    add_paragraph(doc, "• Slide generation")
    add_paragraph(doc, "• Image generation with Gemini")
    add_paragraph(doc, "• Document analysis")
    
    add_heading(doc, "Socket Manager (socket_manager.py)", level=3)
    add_paragraph(doc, "Real-time communication:")
    add_paragraph(doc, "• WebSocket connections")
    add_paragraph(doc, "• Room-based messaging")
    add_paragraph(doc, "• Event broadcasting")
    add_paragraph(doc, "• Connection management")
    
    doc.add_page_break()

def add_frontend_docs(doc):
    """Add frontend documentation"""
    add_heading(doc, "6. Frontend Documentation", level=1)
    
    add_heading(doc, "6.1 Project Structure", level=2)
    add_code_block(doc, """frontend/src/
├── components/          # Reusable components
│   ├── Button.js       # Standardized button component
│   ├── Layout.js       # Main layout wrapper
│   ├── LoadingSpinner.js
│   ├── ConfirmModal.js
│   ├── ChatInterface.js
│   └── ...
├── pages/              # Page components
│   ├── DashboardPage.js
│   ├── NotesPage.js
│   ├── DocumentSessionPage.js
│   ├── ClassroomPage.js
│   ├── YouTubeSummarizerPage.js
│   └── ...
├── contexts/           # React contexts
│   ├── AuthContext.js
│   └── SocketContext.js
├── utils/
│   ├── api.js          # API client
│   └── ...
├── App.js              # Root component
└── index.js            # Entry point""")
    
    add_heading(doc, "6.2 Key Pages", level=2)
    
    pages = {
        "Dashboard": "Main landing page showing classrooms and recent activity",
        "Notes": "Document management with upload, create, and AI sessions",
        "Document Session": "AI-powered document analysis and chat",
        "YouTube Summarizer": "Video processing and study material generation",
        "Classroom": "Virtual classroom with channels and real-time chat",
        "Friends": "Friend management and requests",
        "Profile": "User settings and profile"
    }
    
    for page, desc in pages.items():
        add_paragraph(doc, f"• {page}: {desc}", bold=True)
    
    add_heading(doc, "6.3 State Management", level=2)
    add_paragraph(doc, "The application uses React Query for server state management:")
    add_paragraph(doc, "• Automatic caching and refetching")
    add_paragraph(doc, "• Optimistic updates")
    add_paragraph(doc, "• Query invalidation")
    add_paragraph(doc, "• Loading and error states")
    
    doc.add_page_break()

def add_database_schema(doc):
    """Add database schema section"""
    add_heading(doc, "7. Database Schema", level=1)
    
    collections = {
        "users": [
            "username: string (unique)",
            "email: string (unique)",
            "hashed_password: string",
            "name: string",
            "friends: array of user IDs",
            "friend_requests: array of user IDs",
            "created_at: datetime",
            "updated_at: datetime"
        ],
        "classrooms": [
            "name: string",
            "description: string",
            "admin_id: ObjectId",
            "members: array of user IDs",
            "invite_code: string (unique)",
            "rooms: array of room documents",
            "created_at: datetime",
            "updated_at: datetime"
        ],
        "messages": [
            "room_id: ObjectId",
            "sender_id: ObjectId",
            "content: string",
            "deleted: boolean",
            "edited: boolean",
            "created_at: datetime",
            "updated_at: datetime"
        ],
        "documents": [
            "user_id: ObjectId",
            "title: string",
            "content: string",
            "file_path: string",
            "file_size: number",
            "status: string",
            "created_at: datetime",
            "updated_at: datetime"
        ],
        "document_sessions": [
            "document_id: ObjectId",
            "user_id: ObjectId",
            "short_summary: string",
            "detailed_summary: string",
            "chat_history: array",
            "flashcards: array",
            "created_at: datetime",
            "updated_at: datetime"
        ],
        "youtube_sessions": [
            "user_id: ObjectId",
            "video_id: string",
            "title: string",
            "thumbnail_url: string",
            "short_summary: string",
            "detailed_summary: string",
            "flashcards: array",
            "slides: array",
            "related_videos: array",
            "created_at: datetime",
            "updated_at: datetime"
        ]
    }
    
    for collection, fields in collections.items():
        add_heading(doc, collection, level=2)
        for field in fields:
            add_paragraph(doc, f"• {field}")
    
    doc.add_page_break()

def add_api_docs(doc):
    """Add API documentation"""
    add_heading(doc, "8. API Documentation", level=1)
    
    add_heading(doc, "8.1 Authentication Endpoints", level=2)
    api_routes = [
        ("POST", "/api/auth/register", "Register new user"),
        ("POST", "/api/auth/login", "User login"),
        ("POST", "/api/auth/refresh", "Refresh JWT token"),
        ("GET", "/api/auth/me", "Get current user")
    ]
    
    for method, endpoint, desc in api_routes:
        add_paragraph(doc, f"{method} {endpoint} - {desc}")
    
    add_heading(doc, "8.2 Document Endpoints", level=2)
    doc_routes = [
        ("GET", "/api/documents", "List all user documents"),
        ("POST", "/api/documents", "Create new document"),
        ("POST", "/api/documents/upload", "Upload document file"),
        ("GET", "/api/documents/{id}", "Get document details"),
        ("PUT", "/api/documents/{id}", "Update document"),
        ("DELETE", "/api/documents/{id}", "Delete document"),
        ("POST", "/api/documents/{id}/session", "Create AI session"),
        ("POST", "/api/documents/session/{id}/chat", "Chat with AI"),
        ("POST", "/api/documents/session/{id}/generate-notes", "Generate notes")
    ]
    
    for method, endpoint, desc in doc_routes:
        add_paragraph(doc, f"{method} {endpoint} - {desc}")
    
    add_heading(doc, "8.3 YouTube Endpoints", level=2)
    yt_routes = [
        ("POST", "/api/youtube/process", "Process YouTube video"),
        ("GET", "/api/youtube/sessions", "Get user sessions"),
        ("POST", "/api/youtube/flashcards", "Generate flashcards"),
        ("POST", "/api/youtube/slides", "Generate slides"),
        ("POST", "/api/youtube/related", "Get related videos")
    ]
    
    for method, endpoint, desc in yt_routes:
        add_paragraph(doc, f"{method} {endpoint} - {desc}")
    
    add_heading(doc, "8.4 Classroom Endpoints", level=2)
    class_routes = [
        ("GET", "/api/classrooms", "List user classrooms"),
        ("POST", "/api/classrooms", "Create classroom"),
        ("GET", "/api/classrooms/{id}", "Get classroom details"),
        ("PUT", "/api/classrooms/{id}", "Update classroom"),
        ("DELETE", "/api/classrooms/{id}", "Delete classroom"),
        ("POST", "/api/classrooms/{id}/rooms", "Create room"),
        ("DELETE", "/api/classrooms/{id}/rooms/{room_id}", "Delete room"),
        ("POST", "/api/classrooms/{id}/members", "Add member"),
        ("DELETE", "/api/classrooms/{id}/members/{user_id}", "Remove member"),
        ("POST", "/api/classrooms/join", "Join with invite code")
    ]
    
    for method, endpoint, desc in class_routes:
        add_paragraph(doc, f"{method} {endpoint} - {desc}")
    
    doc.add_page_break()

def add_setup_guide(doc):
    """Add setup and installation guide"""
    add_heading(doc, "9. Setup & Installation", level=1)
    
    add_heading(doc, "9.1 Prerequisites", level=2)
    add_paragraph(doc, "• Python 3.12+")
    add_paragraph(doc, "• Node.js 18+")
    add_paragraph(doc, "• MongoDB 6.0+")
    add_paragraph(doc, "• Git")
    
    add_heading(doc, "9.2 Backend Setup", level=2)
    add_paragraph(doc, "Step 1: Clone the repository", bold=True)
    add_code_block(doc, "git clone https://github.com/Mohammad-Sofyan-Abdullah/DataMinors_Edu.git\ncd DataMinors_Edu/backend")
    
    add_paragraph(doc, "Step 2: Create virtual environment", bold=True)
    add_code_block(doc, "python -m venv venv\nsource venv/bin/activate  # On Windows: venv\\Scripts\\activate")
    
    add_paragraph(doc, "Step 3: Install dependencies", bold=True)
    add_code_block(doc, "pip install -r requirements.txt")
    
    add_paragraph(doc, "Step 4: Configure environment variables", bold=True)
    add_paragraph(doc, "Create .env file with:")
    add_code_block(doc, """MONGODB_URL=mongodb://localhost:27017
DATABASE_NAME=dataminors_edu
SECRET_KEY=your-secret-key-here
GROQ_API_KEY=your-groq-api-key
GEMINI_API_KEY=your-gemini-api-key
FRONTEND_URL=http://localhost:3000""")
    
    add_paragraph(doc, "Step 5: Run the server", bold=True)
    add_code_block(doc, "python -m uvicorn app.main:app --reload --host 0.0.0.0 --port 8000")
    
    add_heading(doc, "9.3 Frontend Setup", level=2)
    add_paragraph(doc, "Step 1: Navigate to frontend directory", bold=True)
    add_code_block(doc, "cd ../frontend")
    
    add_paragraph(doc, "Step 2: Install dependencies", bold=True)
    add_code_block(doc, "npm install")
    
    add_paragraph(doc, "Step 3: Configure environment", bold=True)
    add_paragraph(doc, "Create .env file with:")
    add_code_block(doc, "REACT_APP_API_URL=http://localhost:8000")
    
    add_paragraph(doc, "Step 4: Run development server", bold=True)
    add_code_block(doc, "npm start")
    
    add_heading(doc, "9.4 MongoDB Setup", level=2)
    add_paragraph(doc, "Option 1: Local MongoDB", bold=True)
    add_code_block(doc, "mongod --dbpath /path/to/data/db")
    
    add_paragraph(doc, "Option 2: MongoDB Atlas", bold=True)
    add_paragraph(doc, "• Create account at mongodb.com/cloud/atlas")
    add_paragraph(doc, "• Create cluster")
    add_paragraph(doc, "• Get connection string")
    add_paragraph(doc, "• Update MONGODB_URL in .env")
    
    doc.add_page_break()

def add_deployment_guide(doc):
    """Add deployment guide"""
    add_heading(doc, "10. Deployment Guide", level=1)
    
    add_heading(doc, "10.1 Production Build", level=2)
    add_paragraph(doc, "Frontend Production Build", bold=True)
    add_code_block(doc, "cd frontend\nnpm run build")
    
    add_paragraph(doc, "This creates optimized production build in build/ directory")
    
    add_heading(doc, "10.2 Environment Variables", level=2)
    add_paragraph(doc, "Ensure all production environment variables are set:")
    add_paragraph(doc, "• API URLs")
    add_paragraph(doc, "• Database connection strings")
    add_paragraph(doc, "• API keys for external services")
    add_paragraph(doc, "• Secret keys for JWT")
    
    add_heading(doc, "10.3 Hosting Options", level=2)
    
    add_paragraph(doc, "Frontend:", bold=True)
    add_paragraph(doc, "• Vercel")
    add_paragraph(doc, "• Netlify")
    add_paragraph(doc, "• AWS S3 + CloudFront")
    add_paragraph(doc, "• GitHub Pages")
    
    add_paragraph(doc, "Backend:", bold=True)
    add_paragraph(doc, "• AWS EC2")
    add_paragraph(doc, "• Heroku")
    add_paragraph(doc, "• DigitalOcean")
    add_paragraph(doc, "• Google Cloud Run")
    
    add_paragraph(doc, "Database:", bold=True)
    add_paragraph(doc, "• MongoDB Atlas (recommended)")
    add_paragraph(doc, "• Self-hosted MongoDB")
    
    add_heading(doc, "10.4 Production Checklist", level=2)
    add_paragraph(doc, "☐ Configure CORS properly")
    add_paragraph(doc, "☐ Enable HTTPS")
    add_paragraph(doc, "☐ Set up monitoring")
    add_paragraph(doc, "☐ Configure backups")
    add_paragraph(doc, "☐ Set up logging")
    add_paragraph(doc, "☐ Performance optimization")
    add_paragraph(doc, "☐ Security audit")

def generate_documentation():
    """Main function to generate complete documentation"""
    print("Generating DataMinors_Edu documentation...")
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "DataMinors_Edu Documentation"
    doc.core_properties.author = "DataMinors_Edu Team"
    doc.core_properties.subject = "Complete Project Documentation"
    
    # Generate sections
    print("Creating title page...")
    create_title_page(doc)
    
    print("Adding table of contents...")
    add_table_of_contents(doc)
    
    print("Adding project overview...")
    add_project_overview(doc)
    
    print("Adding architecture...")
    add_architecture(doc)
    
    print("Adding technology stack...")
    add_technology_stack(doc)
    
    print("Adding features...")
    add_features(doc)
    
    print("Adding backend documentation...")
    add_backend_docs(doc)
    
    print("Adding frontend documentation...")
    add_frontend_docs(doc)
    
    print("Adding database schema...")
    add_database_schema(doc)
    
    print("Adding API documentation...")
    add_api_docs(doc)
    
    print("Adding setup guide...")
    add_setup_guide(doc)
    
    print("Adding deployment guide...")
    add_deployment_guide(doc)
    
    # Save document
    output_file = "DataMinors_Edu_Documentation.docx"
    doc.save(output_file)
    print(f"\n✅ Documentation generated successfully: {output_file}")
    print(f"📄 File size: {os.path.getsize(output_file) / 1024:.2f} KB")

if __name__ == "__main__":
    generate_documentation()
